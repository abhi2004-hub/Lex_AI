"""
Document Processing Pipeline for LexAI Platform
Handles document parsing, preprocessing, and text extraction
"""

import os
import PyPDF2
import docx
from typing import Dict, List, Optional
import logging
import re
from datetime import datetime
import hashlib

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing and text extraction"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        
    def process_document(self, file_path: str, original_filename: str) -> Dict:
        """
        Process uploaded document and extract metadata
        
        Args:
            file_path: Path to the uploaded file
            original_filename: Original filename from upload
            
        Returns:
            Dictionary containing document metadata and extracted text
        """
        try:
            # Validate file
            if not self._validate_file(file_path, original_filename):
                raise ValueError("Invalid file format or size")
            
            # Extract file metadata
            file_metadata = self._extract_file_metadata(file_path, original_filename)
            
            # Extract text based on file type
            file_ext = os.path.splitext(original_filename)[1].lower()
            
            if file_ext == '.pdf':
                extracted_data = self._process_pdf(file_path)
            elif file_ext == '.docx':
                extracted_data = self._process_docx(file_path)
            elif file_ext == '.txt':
                extracted_data = self._process_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Combine metadata and extracted data
            document_data = {
                **file_metadata,
                **extracted_data,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            # Clean and preprocess text
            document_data['cleaned_text'] = self._preprocess_text(document_data['text'])
            
            # Calculate additional metadata
            document_data['word_count'] = len(document_data['cleaned_text'].split())
            document_data['char_count'] = len(document_data['cleaned_text'])
            document_data['line_count'] = len(document_data['cleaned_text'].split('\n'))
            
            logger.info(f"Successfully processed document: {original_filename}")
            return document_data
            
        except Exception as e:
            logger.error(f"Error processing document {original_filename}: {str(e)}")
            raise
    
    def _validate_file(self, file_path: str, filename: str) -> bool:
        """Validate file format and size"""
        # Check file extension
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in self.supported_formats:
            logger.error(f"Unsupported file format: {file_ext}")
            return False
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            logger.error(f"File too large: {file_size} bytes")
            return False
        
        # Check if file exists and is readable
        if not os.path.exists(file_path) or not os.access(file_path, os.R_OK):
            logger.error(f"File not accessible: {file_path}")
            return False
        
        return True
    
    def _extract_file_metadata(self, file_path: str, original_filename: str) -> Dict:
        """Extract basic file metadata"""
        stat = os.stat(file_path)
        
        # Generate file hash for integrity checking
        file_hash = self._generate_file_hash(file_path)
        
        metadata = {
            'original_filename': original_filename,
            'file_size': stat.st_size,
            'file_type': os.path.splitext(original_filename)[1].lower(),
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'file_hash': file_hash
        }
        
        return metadata
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF document and extract text"""
        try:
            text_content = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append({
                                'page_number': page_num + 1,
                                'text': page_text.strip()
                            })
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            # Combine all page text
            full_text = '\n\n'.join([page['text'] for page in text_content])
            
            return {
                'text': full_text,
                'page_count': page_count,
                'pages': text_content,
                'format': 'PDF'
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict:
        """Process DOCX document and extract text"""
        try:
            doc = docx.Document(file_path)
            
            paragraphs = []
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraph_text = para.text.strip()
                    paragraphs.append({
                        'paragraph_number': len(paragraphs) + 1,
                        'text': paragraph_text,
                        'style': para.style.name if para.style else 'Normal'
                    })
                    full_text.append(paragraph_text)
            
            # Extract tables if any
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Combine all text
            combined_text = '\n\n'.join(full_text)
            
            # Estimate page count (rough approximation)
            page_count = max(1, len(combined_text) // 3000)  # Assume ~3000 chars per page
            
            return {
                'text': combined_text,
                'page_count': page_count,
                'paragraphs': paragraphs,
                'tables': tables,
                'format': 'DOCX'
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise
    
    def _process_txt(self, file_path: str) -> Dict:
        """Process TXT document and extract text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Split into paragraphs
            paragraphs = []
            paragraph_text = text_content.split('\n\n')
            
            for i, para in enumerate(paragraph_text):
                if para.strip():
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'text': para.strip()
                    })
            
            # Estimate page count
            page_count = max(1, len(text_content) // 3000)
            
            return {
                'text': text_content,
                'page_count': page_count,
                'paragraphs': paragraphs,
                'format': 'TXT'
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                
                page_count = max(1, len(text_content) // 3000)
                
                return {
                    'text': text_content,
                    'page_count': page_count,
                    'format': 'TXT',
                    'encoding': 'latin-1'
                }
            except Exception as e:
                logger.error(f"Error reading TXT file with alternative encoding {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error processing TXT file {file_path}: {str(e)}")
            raise
    
    def _preprocess_text(self, text: str) -> str:
        """
        Preprocess extracted text for better AI analysis
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and preprocessed text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\"\'\/\n]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple line breaks to double
        text = re.sub(r'\n', ' ', text)  # Single line breaks to space
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?)])', r'\1', text)
        text = re.sub(r'([(])\s+', r'\1', text)
        
        # Remove page numbers and headers/footers (common patterns)
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)  # Page numbers
        text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text)  # Page X of Y
        text = re.sub(r'^.{0,50}(\n|$)', '', text, flags=re.MULTILINE)  # Potential headers
        
        # Clean up extra spaces
        text = ' '.join(text.split())
        
        return text.strip()
    
    def _generate_file_hash(self, file_path: str) -> str:
        """Generate SHA-256 hash of file for integrity checking"""
        hash_sha256 = hashlib.sha256()
        
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        
        return hash_sha256.hexdigest()
    
    def detect_language(self, text: str) -> str:
        """
        Detect the language of the document text
        
        Args:
            text: Document text to analyze
            
        Returns:
            ISO 639-1 language code
        """
        try:
            # Simple language detection based on common words
            # In production, would use libraries like langdetect or polyglot
            
            text_sample = text[:1000].lower()  # Use first 1000 characters
            
            # Common language indicators
            language_patterns = {
                'en': ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for'],
                'es': ['el', 'la', 'y', 'o', 'pero', 'en', 'de', 'para', 'por'],
                'fr': ['le', 'la', 'et', 'ou', 'mais', 'dans', 'de', 'pour', 'par'],
                'de': ['der', 'die', 'das', 'und', 'oder', 'aber', 'in', 'zu', 'für'],
                'it': ['il', 'la', 'e', 'o', 'ma', 'in', 'di', 'per', 'da']
            }
            
            language_scores = {}
            
            for lang, patterns in language_patterns.items():
                score = sum(1 for pattern in patterns if pattern in text_sample)
                language_scores[lang] = score
            
            # Return language with highest score, default to English
            best_language = max(language_scores, key=language_scores.get)
            
            # If no strong indicators, default to English
            if language_scores[best_language] < 2:
                return 'en'
            
            return best_language
            
        except Exception as e:
            logger.error(f"Error detecting language: {str(e)}")
            return 'en'  # Default to English
    
    def extract_document_structure(self, text: str) -> Dict:
        """
        Extract document structure (headings, sections, etc.)
        
        Args:
            text: Preprocessed document text
            
        Returns:
            Dictionary containing document structure information
        """
        try:
            structure = {
                'headings': [],
                'sections': [],
                'numbered_lists': [],
                'bullet_points': []
            }
            
            lines = text.split('\n')
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Detect headings (all caps, or ending with colon)
                if line.isupper() or line.endswith(':'):
                    structure['headings'].append({
                        'line_number': i + 1,
                        'text': line,
                        'level': 1 if line.isupper() else 2
                    })
                
                # Detect numbered sections
                numbered_match = re.match(r'^(\d+\.?\d*)\s+(.+)', line)
                if numbered_match:
                    structure['sections'].append({
                        'line_number': i + 1,
                        'number': numbered_match.group(1),
                        'title': numbered_match.group(2)
                    })
                
                # Detect numbered lists
                list_match = re.match(r'^([a-zA-Z]|\d+)\)\s+(.+)', line)
                if list_match:
                    structure['numbered_lists'].append({
                        'line_number': i + 1,
                        'marker': list_match.group(1),
                        'text': list_match.group(2)
                    })
                
                # Detect bullet points
                bullet_match = re.match(r'^[•\-\*]\s+(.+)', line)
                if bullet_match:
                    structure['bullet_points'].append({
                        'line_number': i + 1,
                        'text': bullet_match.group(1)
                    })
            
            return structure
            
        except Exception as e:
            logger.error(f"Error extracting document structure: {str(e)}")
            return {'headings': [], 'sections': [], 'numbered_lists': [], 'bullet_points': []}
    
    def split_into_clauses(self, text: str, max_clause_length: int = 500) -> List[str]:
        """
        Split document text into potential clauses
        
        Args:
            text: Preprocessed document text
            max_clause_length: Maximum length for each clause
            
        Returns:
            List of clause texts
        """
        try:
            clauses = []
            
            # Split by sentence endings first
            sentences = re.split(r'[.!?]+', text)
            
            current_clause = ""
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                # If adding this sentence would exceed max length, start new clause
                if len(current_clause) + len(sentence) > max_clause_length and current_clause:
                    clauses.append(current_clause.strip())
                    current_clause = sentence
                else:
                    current_clause += ". " + sentence if current_clause else sentence
            
            # Add the last clause
            if current_clause.strip():
                clauses.append(current_clause.strip())
            
            # Filter out very short clauses
            clauses = [clause for clause in clauses if len(clause) > 20]
            
            return clauses
            
        except Exception as e:
            logger.error(f"Error splitting into clauses: {str(e)}")
            return [text]  # Return full text as single clause if splitting fails
